"""
Core converter module for parsing Moodle HTML quiz to text and Word format.
Uses BeautifulSoup4 for HTML parsing and python-docx for document generation.
"""

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Tuple, Dict
import re


class QuizParser:
    """Parser for extracting quiz questions from HTML."""
    
    @staticmethod
    def extract_text_from_element(element):
        """
        Extract cleaned text from a BeautifulSoup element.
        Handles nested elements and removes hidden text.
        """
        if not element:
            return ""
        
        # Create a copy to avoid modifying original
        text_parts = []
        
        for child in element.descendants:
            # Skip elements with 'accesshide' class (hidden content)
            if hasattr(child.parent, 'get'):
                classes = child.parent.get('class', [])
                if 'accesshide' in classes:
                    continue
            
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    text_parts.append(text)
        
        # Join and clean up
        full_text = ' '.join(text_parts)
        # Remove extra whitespace
        full_text = re.sub(r'\s+', ' ', full_text)
        return full_text.strip()
    
    @staticmethod
    def parse_questions(html_content: str) -> List[Dict]:
        """
        Parse Moodle quiz HTML and extract questions with answers.
        
        Args:
            html_content: HTML string from Moodle quiz page
            
        Returns:
            List of question dictionaries with structure:
            {
                'question_number': int,
                'question_text': str,
                'answers': [
                    {'letter': 'A', 'content': str},
                    {'letter': 'B', 'content': str},
                    ...
                ]
            }
        """
        # Use html.parser to avoid lxml DLL issues
        soup = BeautifulSoup(html_content, 'html.parser')
        questions = []
        
        # Find all question divs: class contains 'que multichoice'
        question_divs = soup.find_all(
            'div',
            class_=lambda c: c and 'que' in c and 'multichoice' in c
        )
        
        for question_num, que_div in enumerate(question_divs, 1):
            try:
                # Extract question text from qtext div
                qtext_div = que_div.find('div', class_='qtext')
                if not qtext_div:
                    continue
                
                question_text = QuizParser.extract_text_from_element(qtext_div)
                if not question_text:
                    continue
                
                # Extract answers
                answers = []
                answer_div = que_div.find('div', class_='answer')
                
                if answer_div:
                    # Find all answer blocks
                    answer_blocks = answer_div.find_all('div', recursive=False)
                    
                    for answer_block in answer_blocks:
                        # Get answer letter from answernumber span
                        answer_number_span = answer_block.find(
                            'span',
                            class_='answernumber'
                        )
                        if not answer_number_span:
                            continue
                        
                        answer_letter = QuizParser.extract_text_from_element(
                            answer_number_span
                        ).strip()
                        
                        # Get answer content from flex-fill div
                        flex_div = answer_block.find('div', class_='flex-fill')
                        answer_content = QuizParser.extract_text_from_element(
                            flex_div
                        ) if flex_div else ""
                        
                        if answer_letter and answer_content:
                            answers.append({
                                'letter': answer_letter,
                                'content': answer_content
                            })
                
                # Only add question if it has valid answers
                if answers:
                    questions.append({
                        'question_number': question_num,
                        'question_text': question_text,
                        'answers': answers
                    })
            
            except Exception as e:
                print(f"Error parsing question {question_num}: {str(e)}")
                continue
        
        return questions


class QuizConverter:
    """Converter for generating text and Word document outputs."""
    
    @staticmethod
    def to_plain_text(questions: List[Dict]) -> str:
        """
        Convert parsed questions to plain text format.
        
        Format:
        Câu 1: [Question text]
        A. [Answer A]
        B. [Answer B]
        ...
        """
        lines = []
        
        for question in questions:
            q_num = question['question_number']
            q_text = question['question_text']
            lines.append(f"Câu {q_num}: {q_text}")
            
            for answer in question['answers']:
                letter = answer['letter']
                content = answer['content']
                lines.append(f"{letter}. {content}")
            
            lines.append("")  # Blank line between questions
        
        return '\n'.join(lines)
    
    @staticmethod
    def to_word_document(questions: List[Dict]) -> Document:
        """
        Convert parsed questions to a formatted Word document.
        
        Args:
            questions: List of parsed question dictionaries
            
        Returns:
            python-docx Document object
        """
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title_run = title.add_run("BÀI TRẮC NGHIỆM")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Add questions
        for question in questions:
            q_num = question['question_number']
            q_text = question['question_text']
            
            # Add question number and text (bold number)
            q_paragraph = doc.add_paragraph(style='List Bullet')
            q_run = q_paragraph.add_run(f"Câu {q_num}: ")
            q_run.font.bold = True
            q_run.font.size = Pt(11)
            q_paragraph.add_run(q_text).font.size = Pt(11)
            
            # Add answers
            for answer in question['answers']:
                letter = answer['letter']
                content = answer['content']
                a_paragraph = doc.add_paragraph(
                    f"{letter}. {content}",
                    style='List Bullet 2'
                )
                a_paragraph.paragraph_format.left_indent = Inches(0.5)
            
            # Add space between questions
            doc.add_paragraph()
        
        return doc


def parse_html_to_text_and_doc(
    html_content: str
) -> Tuple[str, Document, int]:
    """
    Main function: Parse HTML and generate both text and Word outputs.
    
    Args:
        html_content: Raw HTML from Moodle quiz
        
    Returns:
        Tuple of (plain_text, document, question_count)
    """
    # Parse HTML
    parser = QuizParser()
    questions = parser.parse_questions(html_content)
    
    # Generate outputs
    converter = QuizConverter()
    text_output = converter.to_plain_text(questions)
    doc_output = converter.to_word_document(questions)
    question_count = len(questions)
    
    return text_output, doc_output, question_count


def export_marked_text(questions: List[Dict], predictions: Dict[int, str]) -> str:
    """
    Export quiz as text with AI-predicted answers marked with checkmarks.
    
    Args:
        questions: List of question dictionaries
        predictions: Dictionary mapping question numbers to predicted answers
        
    Returns:
        Plain text with marked answers
    """
    converter = QuizConverter()
    output = []
    
    for q in questions:
        q_num = q.get('question_number', q.get('number', 0))
        q_text = q.get('question_text', q.get('text', ''))
        
        output.append(f"\n{'='*60}")
        output.append(f"Câu {q_num}: {q_text}")
        output.append('='*60)
        
        predicted = predictions.get(q_num, 'A')
        
        for ans in q.get('answers', []):
            letter = ans.get('letter', '')
            content = ans.get('content', '')
            
            # Mark with checkmark if this is the predicted answer
            mark = "✓ " if letter == predicted else "  "
            output.append(f"{mark}{letter}. {content}")
        
    output.append(f"\n{'='*60}")
    output.append("Kết quả được AI phân tích")
    output.append('='*60)
    
    return '\n'.join(output)


def export_marked_docx(questions: List[Dict], predictions: Dict[int, str]) -> Document:
    """
    Export quiz as Word document with AI-predicted answers highlighted in green.
    
    Args:
        questions: List of question dictionaries
        predictions: Dictionary mapping question numbers to predicted answers
        
    Returns:
        Word document with marked answers
    """
    from docx.shared import RGBColor
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Bài Trắc Nghiệm - Kết Quả Phân Tích AI', level=1)
    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for q in questions:
        q_num = q.get('question_number', q.get('number', 0))
        q_text = q.get('question_text', q.get('text', ''))
        
        predicted = predictions.get(q_num, 'A')
        
        # Add question with AI answer annotation
        q_heading = doc.add_heading(f'Câu {q_num}', level=3)
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(q_text)
        
        # Add AI answer annotation at the end
        annotation_run = q_para.add_run(f"  // {predicted} (đáp án đúng)")
        annotation_run.font.italic = True
        annotation_run.font.color.rgb = RGBColor(0, 128, 0)
        annotation_run.font.size = Pt(10)
        
        # Add answers with highlighting for predicted answer
        for ans in q.get('answers', []):
            letter = ans.get('letter', '')
            content = ans.get('content', '')
            
            # Create paragraph for answer
            para = doc.add_paragraph(style='List Bullet')
            
            # Add letter and content
            run = para.add_run(f"{letter}. {content}")
            
            is_predicted = letter == predicted
            if is_predicted:
                # Highlight predicted answer in green
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                # Add checkmark
                checkmark = para.add_run(" ✓")
                checkmark.font.bold = True
                checkmark.font.color.rgb = RGBColor(0, 128, 0)
        
        # Add spacing
        doc.add_paragraph()
    
    return doc


def export_unmarked_docx(questions: List[Dict]) -> Document:
    """
    Export quiz as Word document without any answers shown.
    
    Args:
        questions: List of question dictionaries
        
    Returns:
        Word document without answers
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Bài Trắc Nghiệm', level=1)
    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for q in questions:
        q_num = q.get('question_number', q.get('number', 0))
        q_text = q.get('question_text', q.get('text', ''))
        
        # Add question only, no answers
        q_heading = doc.add_heading(f'Câu {q_num}', level=3)
        q_para = doc.add_paragraph(q_text)
        
        # Add blank space for written answer
        blank = doc.add_paragraph()
        blank.paragraph_format.left_indent = Inches(0.5)
        blank.text = "_" * 40
        
        # Add spacing
        doc.add_paragraph()
    
    return doc

